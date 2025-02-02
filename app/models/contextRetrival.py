import pinecone
from sentence_transformers import SentenceTransformer, util
from typing import List, Dict, Union
import os
from dotenv import load_dotenv
import io
from werkzeug.datastructures import FileStorage
import re
import requests
import numpy as np
import tempfile
import PyPDF2
import docx
import numpy as np

load_dotenv()

class ContextRetriever:
    def __init__(self):
        # Get API key from environment variables
        api_key = os.getenv('PINECONE_API_KEY')
        index_name = os.getenv('PINECONE_INDEX_NAME')
        self.upload_dir = "app/data/upload"
        if not api_key:
            raise ValueError("PINECONE_API_KEY not found in environment variables.")

        # Initialize Pinecone
        self.pc = pinecone.Pinecone(api_key=api_key)
        self.index_name = index_name

        # Initialize the embedding model
        self.model = SentenceTransformer('all-MiniLM-L6-v2')

        # Get the existing index
        if index_name not in self.pc.list_indexes().names():
            raise ValueError(f"Index '{index_name}' does not exist. Please initialize it first.")

        self.index = self.pc.Index(index_name)

        self.chunk_size = 5
        self.chunk_overlap = 2

    def retrieve_context(self, topic: str, top_k: int = 5) -> List[Dict]:
        """Retrieve context based on the topic passed as an argument."""
        try:
            # Generate embedding for the topic
            topic_embedding = self.model.encode(topic)

            # Query Pinecone index
            results = self.index.query(
                vector=topic_embedding.tolist(),
                top_k=top_k,
                include_metadata=True
            )

            # Format results
            formatted_results = []
            for match in results['matches']:
                formatted_results.append({
                    'text': match['metadata']['text'],
                    'source': match['metadata']['filename'],
                    'score': match['score']
                })

            return formatted_results

        except Exception as e:
            print(f"Error during context retrieval: {str(e)}")
            return []
        
    def _split_into_sentences(self, text: str) -> List[str]:
        """
        Split text into sentences using regex pattern matching.
        Handles common sentence endings (., !, ?) while accounting for common abbreviations.
        """
        # Clean the text
        text = text.strip()
        
        # Split on sentence endings followed by capital letters
        # This pattern looks for ., !, or ? followed by spaces and a capital letter
        sentences = re.split(r'(?<=[.!?])\s+(?=[A-Z])', text)
        
        # Further clean the sentences
        sentences = [s.strip() for s in sentences if s.strip()]
        return sentences

    def _create_chunks(self, text: str) -> List[str]:
        """
        Create overlapping chunks from the input text.
        Each chunk contains a specified number of sentences with overlap.
        """
        # Split text into sentences
        sentences = self._split_into_sentences(text)
        chunks = []
        
        # Create chunks with overlap
        for i in range(0, len(sentences), self.chunk_size - self.chunk_overlap):
            chunk = sentences[i:i + self.chunk_size]
            if chunk:  # Only add non-empty chunks
                chunks.append(" ".join(chunk))
        
        return chunks

    def _process_file(self, file: Union[FileStorage, io.BytesIO]) -> str:
        """
        Process the uploaded file and return its content as text.
        """
        if isinstance(file, FileStorage):
            content = file.read()
        else:
            content = file.read()
            
        # Try different encodings if UTF-8 fails
        encodings = ['utf-8', 'latin-1', 'cp1252']
        for encoding in encodings:
            try:
                return content.decode(encoding)
            except UnicodeDecodeError:
                continue
        
        raise ValueError("Unable to decode file content with supported encodings")

    # def upload_context(self, files: List[Union[FileStorage, io.BytesIO]], filenames: List[str]):
    def upload_context(self):
        try:
            total_chunks = 0
            chunk_names = []  # List to store the names of the chunks
            batch_size = 100
            vectors_batch = []
            upload_dir = "app/data/upload"

            files = [f for f in os.listdir(upload_dir) if os.path.isfile(os.path.join(upload_dir, f))]

            for filename in files:
                filepath = os.path.join(upload_dir, filename)
                clean_filename = os.path.splitext(filename)[0].replace('~$', '')

                try:
                    # Extract text based on file type
                    text_content = ""
                    if filename.lower().endswith('.pdf'):
                        from PyPDF2 import PdfReader
                        reader = PdfReader(filepath)
                        text_content = "\n".join([page.extract_text() for page in reader.pages])
                    elif filename.lower().endswith('.docx'):
                        from docx import Document
                        doc = Document(filepath)
                        text_content = "\n".join([para.text for para in doc.paragraphs])
                    else:
                        with open(filepath, 'r', encoding='utf-8', errors='replace') as file:
                            text_content = file.read()

                    # Process text content into chunks and embeddings
                    chunks = self._create_chunks(text_content)
                    chunk_embeddings = self.model.encode(chunks)

                    for i, (chunk, embedding) in enumerate(zip(chunks, chunk_embeddings)):
                        vector_id = f"{clean_filename}_{i}"
                        chunk_name = f"{clean_filename}_chunk_{i}"
                        chunk_names.append(chunk_name)

                        vectors_batch.append({
                            'id': vector_id,
                            'values': embedding.tolist(),
                            'metadata': {
                                'text': chunk,
                                'filename': clean_filename,
                                'chunk_index': i
                            }
                        })

                        if len(vectors_batch) >= batch_size:
                            self.index.upsert(vectors=vectors_batch)
                            total_chunks += len(vectors_batch)
                            vectors_batch = []

                except Exception as e:
                    print(f"Error processing {filename}: {str(e)}")
                    continue

            if vectors_batch:
                self.index.upsert(vectors=vectors_batch)
                total_chunks += len(vectors_batch)

            return {
                'status': 'success',
                'total_chunks': total_chunks,
                'chunk_names': chunk_names,  # Return the names of the chunks
                'message': f"Successfully uploaded {total_chunks} chunks to the index."
            }

        except Exception as e:
            print(f"Error during context upload: {str(e)}")
            raise

    def retrieve_User_Context(self, topic: str, userDocURL: str):
        try:
            # Determine file type
            file_extension = ".pdf" if userDocURL.endswith(".pdf") else ".docx"
            filename = os.path.basename(userDocURL).split("?")[0]  # Remove query params if any
            filepath = os.path.join(self.upload_dir, filename)
            
            # Download and save file to upload directory
            response = requests.get(userDocURL, timeout=10)
            response.raise_for_status()
            
            with open(filepath, "wb") as file:
                file.write(response.content)
                
            # Extract text based on file type
            extracted_text = []
            if file_extension == ".pdf":
                with open(filepath, "rb") as pdf_file:
                    pdf_reader = PyPDF2.PdfReader(pdf_file)
                    for page in pdf_reader.pages:
                        text = page.extract_text()
                        if text:
                            extracted_text.extend(text.split("\n"))
            elif file_extension == ".docx":
                doc = docx.Document(filepath)
                extracted_text = [para.text for para in doc.paragraphs if para.text.strip()]
            else:
                raise ValueError("Unsupported file format. Only PDF and DOCX are supported.")
            
            # Remove empty lines
            extracted_text = [line.strip() for line in extracted_text if line.strip()]
            if not extracted_text:
                raise ValueError("No text extracted from the document.")
            
            # Compute embeddings
            topic_embedding = self.model.encode(topic, convert_to_tensor=True)
            doc_embeddings = self.model.encode(extracted_text, convert_to_tensor=True)
            
            # Compute similarity scores
            similarity_scores = util.pytorch_cos_sim(topic_embedding, doc_embeddings)[0]
            
            # Get top 5 matches
            top_indices = np.argsort(similarity_scores.cpu().numpy())[::-1][:5]
            top_sentences = [extracted_text[i] for i in top_indices]
            
            response = self.upload_context()
            
            # Delete the downloaded file after uploading
            if os.path.exists(filepath):
                os.remove(filepath)

            return {
                "status": "success",
                "top_matches": top_sentences,
                "file_saved_at": filepath
            }
            
        except requests.exceptions.RequestException as e:
            return {"status": "error", "message": f"Error downloading document: {e}"}
        except (PyPDF2.errors.PdfReadError, docx.opc.exceptions.PackageNotFoundError) as e:
            return {"status": "error", "message": f"Error reading document: {e}"}
        except ValueError as e:
            return {"status": "error", "message": f"Processing error: {e}"}
        except Exception as e:
            return {"status": "error", "message": f"Unexpected error: {e}"}
        finally:
            if os.path.exists(filepath):
                os.remove(filepath)

